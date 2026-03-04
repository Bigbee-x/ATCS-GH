#!/usr/bin/env python3
"""Generate the ATCS-GH Project Progress Report as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

doc = Document()

# ── Page Setup ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Custom Styles ─────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 4):
    hstyle = doc.styles[f"Heading {level}"]
    hfont = hstyle.font
    hfont.name = "Calibri"
    hfont.color.rgb = RGBColor(0x1A, 0x56, 0x76)


# ── Helper Functions ──────────────────────────────────────────────────────
def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ATCS-GH")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Adaptive Traffic Control System — Ghana")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Project Progress Report")
run.font.size = Pt(16)
run.bold = True

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
loc = doc.add_paragraph()
loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = loc.add_run("Simulation Target: Achimota/Neoplan Junction, Accra, Ghana")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Project Overview & Vision",
    "2. How It All Started",
    "3. System Architecture",
    "4. Phase 1 — The Baseline (Fixed Timer)",
    "5. Phase 2 — AI Development Journey",
    "6. Expanding the AI — From 2 Actions to 5",
    "7. The Training Challenges",
    "8. Calibrating to Real-World Accra",
    "9. Current State of the Project",
    "10. Where We're Headed — Future Goals",
    "11. Technical Specifications Summary",
    "12. Lessons Learned",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW & VISION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Project Overview & Vision", level=1)

doc.add_paragraph(
    "ATCS-GH (Adaptive Traffic Control System — Ghana) is a research and engineering "
    "project that uses Deep Reinforcement Learning (DRL) to intelligently control traffic "
    "signals at a real Accra junction. The system replaces traditional fixed-timer traffic "
    "lights with an AI agent that observes real-time traffic conditions and dynamically "
    "adjusts signal phases to minimise wait times, reduce congestion, and prioritise "
    "emergency vehicles."
)

doc.add_paragraph(
    "The project is not just a backend simulation — it includes a full real-time 3D "
    "visualiser built in Godot 4 that connects to the AI via WebSocket, rendering live "
    "traffic flow, signal states, congestion overlays, and performance metrics. Every "
    "line of code in the visualiser is procedural — no external assets, no plugins, "
    "purely programmatic 3D graphics."
)

doc.add_heading("The Core Goal", level=2)
doc.add_paragraph(
    "Demonstrate that a DQN-based AI agent can reduce average vehicle wait time by at "
    "least 40% compared to a standard fixed-timer traffic light at the Achimota/Neoplan "
    "Junction on Accra's N6 Nsawam Road — one of the city's most congested corridors."
)

doc.add_heading("Why This Matters", level=2)
doc.add_paragraph(
    "Accra, Ghana's capital, faces severe traffic congestion. Many intersections still "
    "rely on fixed-timer signals that cannot adapt to changing traffic patterns throughout "
    "the day. This project explores whether AI-driven adaptive signals could be a viable "
    "solution for Ghanaian cities, using the Achimota junction as a proof of concept."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 2. HOW IT ALL STARTED
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. How It All Started", level=1)

doc.add_heading("Setting Up the Foundation", level=2)
doc.add_paragraph(
    "The project began with building a traffic simulation using SUMO (Simulation of Urban "
    "MObility), an open-source traffic simulator developed by the German Aerospace Center "
    "(DLR). We defined a 4-way junction with North-South and East-West roads, configured "
    "vehicle routes with morning rush-hour traffic patterns, and set up the TraCI (Traffic "
    "Control Interface) bridge to control SUMO programmatically from Python."
)

doc.add_heading("The First Milestone — A Working Baseline", level=2)
doc.add_paragraph(
    "Before we could build an AI, we needed something to beat. Phase 1 established a "
    "fixed-timer baseline: a simple traffic light that gives 45 seconds of green to "
    "North-South traffic, then 3 seconds of yellow, then 45 seconds to East-West, then "
    "3 seconds of yellow — repeating endlessly. This naive 96-second cycle doesn't "
    "account for traffic volume, time of day, or any real-world complexity."
)

doc.add_paragraph(
    "The baseline simulation ran for 2 simulated hours (7,200 seconds) of morning rush "
    "hour traffic. The initial baseline on the generic network measured an average wait "
    "time of approximately 399 seconds — our first target to beat."
)

doc.add_heading("Building the 3D Visualiser from Scratch", level=2)
doc.add_paragraph(
    "One of the most ambitious aspects of the project was building a real-time 3D "
    "visualiser in Godot 4 Engine. The hard constraint: everything had to be fully "
    "procedural. No 3D model files, no texture images, no plugin addons — every road, "
    "every vehicle, every traffic light, every UI element is generated purely through "
    "GDScript code at runtime."
)

doc.add_paragraph("The visualiser includes:")
add_bullet("Procedurally generated road geometry with lane markings and crosswalks")
add_bullet("3D traffic lights with red/yellow/green signal spheres that change in real-time")
add_bullet("Animated vehicles (cars, trotros, ambulances) that spawn, move, and despawn")
add_bullet("Per-lane congestion heat overlays (green → yellow → red)")
add_bullet("Live performance metrics panel (wait times, queue lengths, throughput)")
add_bullet("Real-time chart plotting of key metrics over time")
add_bullet("Emergency vehicle override buttons (one per approach)")
add_bullet("Audio system with ambient city sounds and emergency sirens")
add_bullet("WebSocket client that receives live data from the Python simulation server")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. System Architecture", level=1)

doc.add_paragraph(
    "ATCS-GH is a multi-component system with three main layers connected via "
    "real-time communication:"
)

doc.add_heading("Architecture Pipeline", level=2)

doc.add_paragraph(
    "SUMO Traffic Simulator  <-->  TraCI Protocol  <-->  Python AI Server (port 8765)  "
    "<-->  WebSocket  <-->  Godot 4 Visualiser"
)

doc.add_heading("Component Breakdown", level=2)

add_table(
    ["Component", "Technology", "Role"],
    [
        ["Traffic Simulation", "SUMO + TraCI", "Microsimulation of vehicles, lanes, traffic lights"],
        ["AI Agent", "PyTorch (Double DQN)", "Observes state, selects signal actions, learns from rewards"],
        ["Training Pipeline", "Python", "Epsilon-greedy exploration, experience replay, target network"],
        ["WebSocket Server", "Python asyncio", "Bridges AI decisions to the visualiser in real-time"],
        ["3D Visualiser", "Godot 4 (GDScript)", "Procedural 3D rendering of live traffic state"],
        ["Metrics & Logging", "Python + CSV", "Records every timestep for analysis and reporting"],
    ],
)

doc.add_paragraph()
doc.add_heading("Key Technical Decisions", level=2)
add_bullet("Double DQN over vanilla DQN — reduces overestimation of Q-values", bold_prefix="Algorithm: ")
add_bullet("Apple Silicon MPS acceleration for neural network training on Mac", bold_prefix="Hardware: ")
add_bullet("setRedYellowGreenState() for direct per-connection signal control (not setPhase())", bold_prefix="Signal Control: ")
add_bullet("All visuals generated in GDScript — zero external dependencies", bold_prefix="Visualiser: ")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 4. PHASE 1 — THE BASELINE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Phase 1 — The Baseline (Fixed Timer)", level=1)

doc.add_paragraph(
    "Every AI project needs a benchmark. Phase 1 established what a 'dumb' traffic "
    "light achieves so we can quantify how much the AI improves things."
)

doc.add_heading("Baseline Configuration", level=2)
add_table(
    ["Parameter", "Value"],
    [
        ["Simulation Duration", "7,200 seconds (2 hours)"],
        ["Cycle", "96 seconds (45s NS green + 3s yellow + 45s EW green + 3s yellow)"],
        ["Timer Type", "Equal split — no adaptation to traffic volume"],
        ["Emergency Handling", "None — ambulances wait like everyone else"],
    ],
)

doc.add_paragraph()
doc.add_heading("Baseline Results (Achimota-Calibrated Network)", level=2)
add_table(
    ["Metric", "Result"],
    [
        ["Vehicles Completed", "5,142"],
        ["Average Wait Time", "81.83 seconds"],
        ["Peak Queue Length", "66 vehicles"],
        ["Average Throughput", "43.3 vehicles/minute"],
        ["Worst Approach (CBD)", "435.88 seconds avg wait"],
        ["Best Approach (Guggisberg)", "18.10 seconds avg wait"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "The massive disparity between approaches (435.88s vs 18.10s) highlights exactly "
    "why fixed timers are inefficient — they give equal green time to roads with vastly "
    "different traffic volumes. The N/S Achimota Forest Road carries 67% of all traffic "
    "but only gets 47% of green time in the equal-split baseline."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 5. PHASE 2 — AI DEVELOPMENT JOURNEY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Phase 2 — AI Development Journey", level=1)

doc.add_heading("The DQN Agent — How It Works", level=2)
doc.add_paragraph(
    "The AI agent uses a Double Deep Q-Network (Double DQN), a reinforcement learning "
    "algorithm where a neural network learns to predict the expected future reward of "
    "each possible action given the current traffic state."
)

doc.add_paragraph("At every decision interval (5 seconds), the agent:")
add_bullet("Observes the current state: queue lengths, vehicle speeds, waiting times, "
           "current phase, emergency vehicle presence")
add_bullet("Selects an action: HOLD current phase, switch to NS_THROUGH, NS_LEFT, "
           "EW_THROUGH, or EW_LEFT")
add_bullet("Receives a reward signal based on how well traffic is flowing")
add_bullet("Stores the experience in a replay buffer for learning")
add_bullet("Updates the neural network weights using batches from the replay buffer")

doc.add_heading("The Reward Function — Teaching the AI What Matters", level=2)
doc.add_paragraph(
    "The reward function is the most critical component. It tells the AI what 'good' "
    "traffic management looks like. Our reward penalises:"
)
add_bullet("Total queue length across all lanes (primary signal)")
add_bullet("Balance penalty — punishes when one approach has much longer queues than "
           "others (weight: 4.0x)")
add_bullet("Maximum wait fairness — extra penalty when any vehicle has waited more "
           "than 30 seconds (weight: 0.4x)")
add_bullet("Emergency vehicle delays — heavy penalty for making ambulances wait")

doc.add_heading("Neural Network Architecture", level=2)
add_table(
    ["Layer", "Details"],
    [
        ["Input", "36-dimensional state vector"],
        ["Hidden 1", "256 neurons, ReLU activation"],
        ["Hidden 2", "256 neurons, ReLU activation"],
        ["Output", "5 neurons (one per action: HOLD, NS_THROUGH, NS_LEFT, EW_THROUGH, EW_LEFT)"],
        ["Optimizer", "Adam (learning rate: 0.0003)"],
        ["Loss", "Huber Loss (SmoothL1)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 6. EXPANDING THE AI — FROM 2 ACTIONS TO 5
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Expanding the AI — From 2 Actions to 5", level=1)

doc.add_heading("The Original Simple System", level=2)
doc.add_paragraph(
    "The initial AI was deliberately simple: 2 actions (NS green or EW green), "
    "17-dimensional state, and basic phase switching via traci.trafficlight.setPhase(). "
    "This worked as a proof of concept but had major limitations:"
)
add_bullet("No protected left turns — left-turning vehicles conflicted with oncoming traffic")
add_bullet("No HOLD action — the agent had to switch every decision interval")
add_bullet("Coarse state representation — approach-level queues, not per-lane")
add_bullet("Simple signal control that didn't match real-world signal timing practices")

doc.add_heading("The Upgrade to 5 Actions / 6 Phases", level=2)
doc.add_paragraph(
    "We completely redesigned the signal control system to be realistic and capable:"
)

add_table(
    ["Aspect", "Before (v1)", "After (v2)"],
    [
        ["Actions", "2 (NS_GREEN, EW_GREEN)", "5 (HOLD, NS_THROUGH, NS_LEFT, EW_THROUGH, EW_LEFT)"],
        ["State Dimensions", "17", "36"],
        ["Signal Phases", "4 (2 green + 2 yellow)", "6 (4 green + 2 yellow)"],
        ["Signal Control", "setPhase() (index-based)", "setRedYellowGreenState() (per-connection)"],
        ["Protected Left Turns", "No", "Yes — dedicated left-turn phases"],
        ["Per-Lane Sensing", "No (approach-level)", "Yes — queue, speed, wait per lane"],
        ["Turn Arrows", "No", "Yes — rendered in 3D visualiser"],
        ["Congestion Overlays", "No", "Yes — per-lane colour-coded heat maps"],
    ],
)

doc.add_paragraph()
doc.add_paragraph(
    "This upgrade required changes across every layer of the system — SUMO configuration, "
    "Python environment, DQN architecture, training pipeline, evaluation scripts, and the "
    "entire Godot visualiser."
)

doc.add_heading("The run_ai.py Discovery", level=2)
doc.add_paragraph(
    "During the Achimota calibration, we discovered a critical bug: the run_ai.py script "
    "(used for live AI-controlled simulation) had never been updated from the old 2-action "
    "system. It was still importing non-existent constants (NS_GREEN, EW_GREEN) and using "
    "setPhase() instead of setRedYellowGreenState(). The same issue existed in "
    "eval_multi_seed.py. Both scripts received complete rewrites to align with the "
    "5-action/6-phase system."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 7. THE TRAINING CHALLENGES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. The Training Challenges", level=1)

doc.add_paragraph(
    "Training an RL agent for traffic control is notoriously difficult. Each training "
    "episode runs a full 2-hour simulation (7,200 timesteps), taking about 2 minutes of "
    "wall-clock time. Here are the key challenges we faced and how we addressed them:"
)

doc.add_heading("Challenge 1: Catastrophic Exploration Failures", level=2)
doc.add_paragraph(
    "During a 300-episode training run, the agent had been steadily improving — reaching "
    "464.9 seconds average wait at episode 266 (close to the then-baseline of 399.1s). "
    "Then episode 299 hit 10,015 seconds, completely destroying the 'last 5 episodes' "
    "average and making it look like the model had failed catastrophically."
)

p = doc.add_paragraph()
run = p.add_run("Root Cause: ")
run.bold = True
p.add_run(
    "EPSILON_MIN was set to 0.05, meaning even after full training, the agent was still "
    "taking ~72 random actions per episode (out of ~1,440 decisions). One badly-timed "
    "random switch during high-traffic could cascade into a gridlock disaster."
)

p = doc.add_paragraph()
run = p.add_run("Fix: ")
run.bold = True
p.add_run("Reduced EPSILON_MIN from 0.05 to 0.01 (only ~14 random actions per episode).")

doc.add_heading("Challenge 2: Gradient Instability", level=2)
doc.add_paragraph(
    "The neural network weights occasionally made large jumps during training, causing "
    "the policy to suddenly forget good behaviours it had learned."
)

p = doc.add_paragraph()
run = p.add_run("Fix: ")
run.bold = True
p.add_run("Tightened gradient clipping from max_norm=10.0 to max_norm=1.0, ensuring "
          "smoother, more stable learning updates.")

doc.add_heading("Challenge 3: Exploration vs Exploitation Balance", level=2)
doc.add_paragraph(
    "The original EPSILON_DECAY of 0.95 caused epsilon to drop too quickly, not giving "
    "the agent enough time to explore the larger 5-action space."
)

p = doc.add_paragraph()
run = p.add_run("Fix: ")
run.bold = True
p.add_run("Adjusted EPSILON_DECAY to 0.975, giving a slower, more gradual transition "
          "from exploration to exploitation.")

doc.add_heading("Challenge 4: Baseline Measurement Confusion", level=2)
doc.add_paragraph(
    "At one point, we questioned whether the 399.1s baseline was even achievable. Analysis "
    "revealed this was from the Phase 1 fixed-timer measurement on the original generic "
    "network — not an Accra-specific target. Once we recalibrated to the Achimota network, "
    "the new baseline measured 81.83s — a much more realistic and meaningful benchmark."
)

doc.add_heading("Challenge 5: SUMO Auto-Generated Phase Mismatch", level=2)
doc.add_paragraph(
    "When we rebuilt the SUMO network for the Achimota calibration, SUMO's netconvert tool "
    "auto-generated a 6-phase traffic light program (with separate protected left-turn "
    "sub-phases). The baseline script was reading these auto-generated phases and assigning "
    "45s to each, resulting in a bloated 186-second cycle and an artificially inflated "
    "baseline of 2,464.98s."
)

p = doc.add_paragraph()
run = p.add_run("Fix: ")
run.bold = True
p.add_run("Rewrote the baseline to use explicit 15-character signal strings for a clean "
          "4-phase program (45+3+45+3 = 96s cycle), bringing the baseline to the correct 81.83s.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 8. CALIBRATING TO REAL-WORLD ACCRA
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Calibrating to Real-World Accra", level=1)

doc.add_heading("The Question That Changed Everything", level=2)
doc.add_paragraph(
    "Partway through training, we asked: 'Can we import actual OpenStreetMap data of "
    "Accra's roads — so we are not simulating a generic city?' This led to extensive "
    "research into real Accra junctions where traffic is a genuine problem."
)

doc.add_heading("Junction Selection Research", level=2)
doc.add_paragraph(
    "We evaluated 11 candidate junctions across Accra, eliminating those that were "
    "incompatible with our 4-arm signalised intersection model:"
)

add_table(
    ["Junction", "Type", "Verdict"],
    [
        ["Kwame Nkrumah Circle", "Multi-level interchange", "Eliminated — interchange, not a signal"],
        ["Tetteh Quarshie", "Cloverleaf interchange", "Eliminated — grade-separated"],
        ["Mallam Junction", "Flyover + roundabout", "Eliminated — multi-level"],
        ["Kaneshie Market", "Complex multi-arm", "Eliminated — too many arms"],
        ["Achimota/Neoplan Junction", "4-arm signalised", "SELECTED"],
        ["37 Military Hospital", "T-junction", "Runner-up (3-arm)"],
        ["Abeka Junction", "T-junction", "Runner-up (3-arm)"],
    ],
)

doc.add_heading("Why Achimota/Neoplan Junction", level=2)
doc.add_paragraph(
    "The Achimota/Neoplan Junction (GPS: 5.6216\u00b0N, 0.2193\u00b0W) emerged as the clear "
    "winner for several reasons:"
)
add_bullet("4-arm signalised intersection — matches our model perfectly")
add_bullet("Located on the N6 Nsawam Road — one of Accra's busiest corridors")
add_bullet("Part of Ghana's ATMC (Automatic Traffic Management & Control) smart signal network")
add_bullet("Heavy N-S commuter traffic with lighter E-W side streets")
add_bullet("Good OpenStreetMap data availability")
add_bullet("Features trotro (minibus) traffic — a distinctly Ghanaian element")

doc.add_heading("Calibration Approach — Option B", level=2)
doc.add_paragraph(
    "Rather than doing a full OSM import (which would require rebuilding the entire "
    "network from scratch), we chose Option B: calibrate our existing model to match "
    "Achimota's real characteristics. This involved:"
)

add_bullet("Renaming all edges to reflect real road names (ACH = Achimota Forest Rd, "
           "AGG = Aggrey Street, GUG = Guggisberg Street)")
add_bullet("Upgrading the East road (Aggrey Street) from 1 lane to 2 lanes at 50 km/h")
add_bullet("Downgrading the West road (Guggisberg Street) to 1 lane at 30 km/h — reflecting "
           "its narrower, lower-capacity nature")
add_bullet("Recalibrating traffic flows: heavy N-S commuter traffic (900 veh/hr northbound), "
           "moderate Aggrey Street (320 veh/hr), light Guggisberg (300 veh/hr)")
add_bullet("Adding trotro (minibus) flows on Aggrey Street and Achimota Forest Road")
add_bullet("Updating all 12+ files across SUMO, Python, and Godot layers")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 9. CURRENT STATE OF THE PROJECT
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Current State of the Project", level=1)

doc.add_heading("What's Working", level=2)
add_bullet("Full SUMO simulation of the Achimota/Neoplan Junction with realistic traffic flows")
add_bullet("5-action/6-phase DQN agent with 36-dimensional state observation")
add_bullet("Complete 3D procedural visualiser in Godot 4 with real-time WebSocket connection")
add_bullet("Proper baseline measurement: 81.83s average wait (target: < 49.10s)")
add_bullet("Training pipeline running on Apple Silicon (MPS acceleration)")
add_bullet("Emergency vehicle preemption system")
add_bullet("Comprehensive metrics logging and CSV export")

doc.add_heading("Current Training Status", level=2)
doc.add_paragraph(
    "A 200-episode training run is currently in progress on the Achimota-calibrated "
    "network. Early results show the agent learning — average wait times are decreasing "
    "from ~1,090s (episode 1, fully random) to the ~290-350s range by episode 90+. "
    "The agent is still far above the 81.83s baseline but is on a downward trajectory."
)

doc.add_heading("Hyperparameter Configuration", level=2)
add_table(
    ["Hyperparameter", "Value", "Notes"],
    [
        ["State Size", "36", "7 lanes x 3 metrics + 4 approach + 6 phase + 1 timer + 4 emergency"],
        ["Action Size", "5", "HOLD, NS_THROUGH, NS_LEFT, EW_THROUGH, EW_LEFT"],
        ["Hidden Layers", "256 x 256", "Two fully-connected layers with ReLU"],
        ["Learning Rate", "0.0003", "Adam optimiser"],
        ["Batch Size", "128", "Sampled from replay buffer"],
        ["Replay Buffer", "100,000", "Experience replay memory"],
        ["Gamma (discount)", "0.99", "High discount for long-term planning"],
        ["Epsilon Start", "1.0", "Fully random initially"],
        ["Epsilon Min", "0.01", "Reduced from 0.05 to prevent catastrophic random actions"],
        ["Epsilon Decay", "0.975", "Gradual transition from exploration to exploitation"],
        ["Gradient Clip", "1.0", "Tightened from 10.0 for training stability"],
        ["Target Net Update", "Every 10 episodes", "Soft update of target network"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 10. WHERE WE'RE HEADED — FUTURE GOALS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Where We're Headed — Future Goals", level=1)

doc.add_heading("Immediate Goals", level=2)
add_bullet("Complete the 200-episode training run and analyse the learning curve")
add_bullet("Achieve the 40% reduction target: average wait < 49.10s (vs 81.83s baseline)")
add_bullet("Run multi-seed evaluation (5 seeds) to verify the trained model's robustness")
add_bullet("Generate a comprehensive evaluation report with statistical confidence")

doc.add_heading("Near-Term Goals", level=2)
add_bullet("Run the --tuned baseline (60s/30s split) to compare AI against a smarter fixed timer")
add_bullet("Optimise the reward function weights based on training analysis")
add_bullet("Test longer training runs (500+ episodes) if the learning curve hasn't plateaued")
add_bullet("Polish the Godot visualiser for demonstration purposes")

doc.add_heading("Long-Term Vision", level=2)
add_bullet("Multi-intersection coordination — extend to adjacent junctions on the N6 corridor")
add_bullet("Time-of-day adaptation — train separate policies for morning rush, midday, and evening")
add_bullet("Full OSM import — build a geographically accurate network from OpenStreetMap data")
add_bullet("Hardware deployment — explore integration with Ghana's ATMC smart signal infrastructure")
add_bullet("Transfer learning — adapt the trained Achimota model to other Accra junctions")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 11. TECHNICAL SPECIFICATIONS SUMMARY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Technical Specifications Summary", level=1)

doc.add_heading("Project File Structure", level=2)
add_table(
    ["Directory", "Contents"],
    [
        ["simulation/", "SUMO network files: edges, nodes, routes, net.xml, sumocfg"],
        ["ai/", "DQN agent (dqn_agent.py), traffic environment (traffic_env.py), model checkpoints"],
        ["scripts/", "Training, baseline, AI runner, WebSocket server, evaluation, metrics logger"],
        ["visualizer/", "Godot 4 project — all procedural GDScript (Main, Intersection, UI, etc.)"],
        ["data/", "CSV results: baseline, training logs, evaluation data"],
    ],
)

doc.add_paragraph()
doc.add_heading("Achimota Junction Network Specifications", level=2)
add_table(
    ["Specification", "Detail"],
    [
        ["Junction Type", "4-arm signalised intersection"],
        ["Location", "Achimota/Neoplan Junction, N6 Nsawam Rd, Accra"],
        ["N/S Road", "Achimota Forest Road — 2 lanes each direction, 50 km/h"],
        ["East Road", "Aggrey Street — 2 lanes, 50 km/h"],
        ["West Road", "Guggisberg Street — 1 lane, 30 km/h"],
        ["Total Incoming Lanes", "7 (N:2, S:2, E:2, W:1)"],
        ["SUMO Connections", "15 turn movements at junction"],
        ["Signal String Length", "15 characters (one per connection)"],
        ["Total Traffic Demand", "~2,270 vehicles/hour (all approaches combined)"],
    ],
)

doc.add_paragraph()
doc.add_heading("State Vector Breakdown (36 dimensions)", level=2)
add_table(
    ["Component", "Dimensions", "Description"],
    [
        ["Per-lane queue", "7", "Normalised queue length for each incoming lane"],
        ["Per-lane speed", "7", "Average speed on each incoming lane"],
        ["Per-lane wait", "7", "Average waiting time per lane"],
        ["Approach queues", "4", "Total queue per approach (N, S, E, W)"],
        ["Phase one-hot", "6", "Current signal phase encoded as binary vector"],
        ["Phase timer", "1", "Time spent in current phase (normalised)"],
        ["Emergency flags", "4", "Emergency vehicle presence per approach"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 12. LESSONS LEARNED
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("12. Lessons Learned", level=1)

doc.add_heading("On Reinforcement Learning", level=2)
add_bullet("Epsilon management is critical — too much exploration late in training "
           "causes catastrophic episodes that undo learning progress")
add_bullet("Gradient clipping matters — without tight clipping, the policy can make "
           "sudden large jumps that erase learned behaviours")
add_bullet("Training for traffic control is inherently slow — each episode is a full "
           "simulation run, so hyperparameter tuning requires patience")
add_bullet("The reward function design is the hardest part — it must balance multiple "
           "competing objectives (throughput, fairness, emergency response)")

doc.add_heading("On System Engineering", level=2)
add_bullet("Keep all components in sync — upgrading the AI from 2 to 5 actions "
           "required changes in every layer, and missed files (run_ai.py) caused "
           "subtle bugs that weren't caught until later")
add_bullet("Explicit signal strings are better than relying on SUMO auto-generation — "
           "auto-generated phases can change unpredictably when the network changes")
add_bullet("Baseline measurements must be re-run whenever the network changes — "
           "the old 399s baseline was meaningless after the Achimota calibration")

doc.add_heading("On Real-World Calibration", level=2)
add_bullet("Not all junctions are suitable — roundabouts, interchanges, and flyovers "
           "don't fit a standard 4-arm signal model")
add_bullet("Asymmetric road widths matter — Guggisberg Street (1 lane) vs Aggrey Street "
           "(2 lanes) creates fundamentally different traffic dynamics")
add_bullet("Local vehicle types (trotros) add realism and complexity to the simulation")
add_bullet("Option B (calibrate existing model) is more practical than Option A (full "
           "OSM import) for initial real-world alignment")

doc.add_paragraph()
doc.add_paragraph()

# ── Footer ────────────────────────────────────────────────────────────────
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("— ATCS-GH Project Progress Report —")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

date_footer = doc.add_paragraph()
date_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_footer.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ══════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "ATCS-GH_Project_Progress_Report.docx"
)
doc.save(output_path)
print(f"Report saved to: {output_path}")
